import os
import wx
import dash_bootstrap_components as dbc
import plotly.express as px
import pandas as pd
import json
import functions
import webbrowser
import time
from collections import OrderedDict
from dash import Dash, html, dcc, Input, Output, dash_table
from docx2pdf import convert
from pdfrw import PdfReader, PdfWriter


class LoggerWriter:
    def __init__(self, level):
        # self.level is really like using log.debug(message)
        # at least in my case
        self.level = level

    def write(self, message):
        # if statement reduces the amount of newlines that are
        # printed to the logger
        if message != '\n':
            self.level(message)

    def flush(self):
        # create a flush method so things can be flushed when
        # the system wants to. Not sure if simply 'printing'
        # sys.stderr is the correct way to do it, but it seemed
        # to work properly for me.
        self.level(sys.stderr)

import logging
import sys
from pathlib import Path
from logging.handlers import TimedRotatingFileHandler
from threading import Timer

# Create Logger if doesn't exist
Path("log").mkdir(parents=True, exist_ok=True)
formatter = logging.Formatter("%(asctime)s - %(levelname)s - %(message)s")
handler = TimedRotatingFileHandler('log/error.log', when="midnight",
interval=1, encoding='utf8')
handler.suffix = "%Y-%m-%d"
handler.setFormatter(formatter)
logger = logging.getLogger()
logger.setLevel(logging.ERROR)
logger.addHandler(handler)
sys.stdout = LoggerWriter(logging.debug)
sys.stderr = LoggerWriter(logging.warning)



port = 8050
def open_browser():
	webbrowser.open_new("http://localhost:{}".format(port))

def get_path(wildcard):
    app = wx.App(None)
    style = wx.FD_OPEN | wx.FD_FILE_MUST_EXIST
    dialog = wx.FileDialog(None, 'SELECIONE O ARQUIVO EXCEL', wildcard=wildcard, style=style)
    if dialog.ShowModal() == wx.ID_OK:
        path = dialog.GetPath()
    else:
        path = None
    dialog.Destroy()
    return path

path=get_path('*.xlsx')
path=path.replace(os.sep, '/')

def get_co(wildcard):
    app = wx.App(None)
    style = wx.FD_OPEN | wx.FD_FILE_MUST_EXIST
    dialog = wx.FileDialog(None, 'SELECIONE O ARQUIVO DE CONTORNO', wildcard=wildcard, style=style)
    if dialog.ShowModal() == wx.ID_OK:
        path = dialog.GetPath()
    else:
        path = None
    dialog.Destroy()
    return path

path_co=get_co('*.shp')
path_co=path_co.replace(os.sep, '/')

def get_pt(wildcard):
    app = wx.App(None)
    style = wx.FD_OPEN | wx.FD_FILE_MUST_EXIST
    dialog = wx.FileDialog(None, 'SELECIONE O ARQUIVO DE PONTOS', wildcard=wildcard, style=style)
    if dialog.ShowModal() == wx.ID_OK:
        path = dialog.GetPath()
    else:
        path = None
    dialog.Destroy()
    return path

path_pt_shp=get_pt('*.shp')
path_pt_shp=path_pt_shp.replace(os.sep, '/')



try:
    resposta_overlap = functions.__count_overlap__(path_co)
except:
    pass
    resposta_overlap = 'NÃO HÁ SOBREPOSIÇÕES'

try:
    pontos_fora = functions._pontos_fora__(path_co, path_pt_shp)

except:
    pass
    pontos_fora = ''
try:
    resposta_header_co = functions.__IDENTIFY_CONTOUR_ERROR__(path_co)

except:
    pass
try:
    estatistica_df = functions.__statistical_module__(path)

except:
    pass

try:
    duplicate_df = functions.__duplicate_df__(path)

except:
    pass

try:
    duplicate_str = functions.__duplicate_str__(duplicate_df)
except:
    pass
    duplicate_str = ''

try:
    functions.__layout_map__(path_co, path_pt_shp)

except:
    pass

try:
    farm = functions.__polygon__(path_co)
    point = functions.__point__(path_pt_shp)
    join_pt = functions.__maps_join_dash__(path, path_pt_shp)
    functions.__export_map_perfil__(path, path_co, path_pt_shp)
    list_png = sorted(functions.__png_perfil_png())
    functions.__graphic_quantification_join_e_amostras_fig(path, path_pt_shp)
    farm_copy = farm.copy(deep=True)
    geocol = farm_copy.pop('geometry')
    farm_copy.insert(0, 'geometry', geocol)
    farm_copy["geometry"] = (farm_copy.to_crs(farm_copy.estimate_utm_crs()).simplify(8).to_crs(farm_copy.crs))
    df = estatistica_df[['Determinação', 'Prof', 'Mín', 'Mean', 'Máx', 'CV%', 'tolerancia']]
    prof_list = df.Prof.unique().tolist()  # GERANDO UMA LISTA DAS PROFUNDIDADE
    without_join_df = functions.___without_join_df__(path, path_co, path_pt_shp)
    list_join_df = without_join_df.prof.unique().tolist()
    join_df = functions.__graphic_quantification_join_e_amostras_df(path, path_pt_shp)
    resposta_erro_text = functions.__resposta_erro_text__(path_co)
    without_join_df.to_excel('download/without_join.xlsx', sheet_name='join', index=False)
except:
    pass
    app = wx.App()
    wx.MessageBox('Erro nos arquivo inseridos, por favor, verificar ', 'Erro', wx.OK | wx.ICON_ERROR)

def __statistic_to_excel__(df):
    # create unique list of names
    UniqueNames = df.Prof.unique()
    DataFrameDict = {elem: pd.DataFrame() for elem in UniqueNames}

    for key in DataFrameDict.keys():
        DataFrameDict[key] = df[:][df.Prof == key]
    writer = pd.ExcelWriter('download/estatistica_determinações.xlsx', engine='xlsxwriter')

    for sheet, frame in DataFrameDict.items():
        frame.to_excel(writer, sheet_name=sheet, index=False)
        workbook = writer.book
        worksheet = writer.sheets[sheet]
        (max_row, max_col) = frame.shape
        column_settings = [{'header': column} for column in frame.columns]
        worksheet.add_table(0, 0, max_row, max_col - 1, {'columns': column_settings, 'style': 'Table Style Light 17'})
        worksheet.set_column(0, max_col - 1, 12)

    writer.save()
try:
    __statistic_to_excel__(df)
except:
    pass
    app = wx.App()
    wx.MessageBox('Erro df ', 'Erro', wx.OK | wx.ICON_ERROR)
def __download_shp__(path, path_pt_shp):
    shp = functions.__shp_export__(path, path_pt_shp)
    # create unique list of names
    UniqueNames = shp.prof.unique()

    # create a data frame dictionary to store your data frames
    DataFrameDict = {elem: pd.DataFrame() for elem in UniqueNames}
    try:
        del DataFrameDict[pd.np.nan]
    except:
        pass
    for key in DataFrameDict.keys():
        DataFrameDict[key] = shp[:][shp.prof == key]

    frame_list = []
    for key, value in DataFrameDict.items():
        frame_list.append(value)

    """
        RENOMEAR COLUNA DE ACORDO COM A PROFUNDIDADE

    """
    shp_list = []
    for i in frame_list:
        if '0 A 20 CM' == i['prof'].head().iloc[0]:
            a20 = i.rename(
                columns={c: c + '1' for c in i.columns if c not in ['ID', 'lab', 'lote', 'prof', 'geometry']})
            # shp_list.append(a20)
        if '0 A 25 CM' == i['prof'].head().iloc[0]:
            a25 = i.rename(
                columns={c: c + '1' for c in i.columns if c not in ['ID', 'lab', 'lote', 'prof', 'geometry']})
            # shp_list.append(a20)

        if '0 A 10 CM' == i['prof'].head().iloc[0]:
            a10 = i.rename(
                columns={c: c + '2' for c in i.columns if c not in ['ID', 'lab', 'lote', 'prof', 'geometry']})
            # shp_list.append(a10)
        if '10 A 20 CM' == i['prof'].head().iloc[0]:
            a10a20 = i.rename(
                columns={c: c + '3' for c in i.columns if c not in ['ID', 'lab', 'lote', 'prof', 'geometry']})
            # shp_list.append(a10a20)
        if '20 A 40 CM' == i['prof'].head().iloc[0]:
            a40 = i.rename(
                columns={c: c + '4' for c in i.columns if c not in ['ID', 'lab', 'lote', 'prof', 'geometry']})

        if '25 A 50 CM' == i['prof'].head().iloc[0]:
            a50 = i.rename(
                columns={c: c + '4' for c in i.columns if c not in ['ID', 'lab', 'lote', 'prof', 'geometry']})
            # shp_list.append(a40)
        if '40 A 60 CM' == i['prof'].head().iloc[0]:
            a60 = i.rename(
                columns={c: c + '5' for c in i.columns if c not in ['ID', 'lab', 'lote', 'prof', 'geometry']})
        # shp_list.append(a60)
        if '60 A 80 CM' == i['prof'].head().iloc[0]:
            a80 = i.rename(
                columns={c: c + '6' for c in i.columns if c not in ['ID', 'lab', 'lote', 'prof', 'geometry']})
            # shp_list.append(a80)
        if '80 A 100 CM' == i['prof'].head().iloc[0]:
            a100 = i.rename(
                columns={c: c + '7' for c in i.columns if c not in ['ID', 'lab', 'lote', 'prof', 'geometry']})

    try:
        shp_list.append(a20)
    except:
        pass

    try:
        shp_list.append(a25)
    except:
        pass

    try:
        shp_list.append(a10)
    except:
        pass

    try:
        shp_list.append(a10a20)
    except:
        pass

    try:
        shp_list.append(a40)
    except:
        pass

    try:
        shp_list.append(a50)
    except:
        pass

    try:
        shp_list.append(a60)
    except:
        pass

    try:
        shp_list.append(a80)
    except:
        pass

    try:
        shp_list.append(a100)
    except:
        pass

    """
        DOWNLOAD DAS RESPECTIVAS PROFUNDIDADES

    """
    schema1 = {'geometry': 'Point',
               'properties': OrderedDict([('ID', 'int'),
                                          ('lab', 'str'),
                                          ('lote', 'str'),
                                          ('prof', 'str'),
                                          ('zn1', 'float:15.2'),
                                          ('mn1', 'float:15.2'),
                                          ('fe1', 'float:15.2'),
                                          ('cu1', 'float:15.2'),
                                          ('b1', 'float:15.2'),
                                          ('s1', 'float:15.2'),
                                          ('sat_al1', 'float:15.2'),
                                          ('al1', 'float:15.2'),
                                          ('p_meh1', 'float:15.2'),
                                          ('p_rem1', 'float:15.2'),
                                          ('p_res1', 'float:15.2'),
                                          ('sat_k1', 'float:15.2'),
                                          ('k1', 'float:15.2'),
                                          ('rel_ca_mg1', 'float:15.2'),
                                          ('sat_mg1', 'float:15.2'),
                                          ('mg1', 'float:15.2'),
                                          ('sat_ca1', 'float:15.2'),
                                          ('ca1', 'float:15.2'),
                                          ('v1', 'float:15.2'),
                                          ('ph1', 'float:15.2'),
                                          ('ctc1', 'float:15.2'),
                                          ('mo1', 'float:15.2'),
                                          ('argila1', 'float:15.2')])}

    schema2 = {'geometry': 'Point',
               'properties': OrderedDict([('ID', 'int'),
                                          ('lab', 'str'),
                                          ('lote', 'str'),
                                          ('prof', 'str'),
                                          ('zn2', 'float:15.2'),
                                          ('mn2', 'float:15.2'),
                                          ('fe2', 'float:15.2'),
                                          ('cu2', 'float:15.2'),
                                          ('b2', 'float:15.2'),
                                          ('s2', 'float:15.2'),
                                          ('sat_al2', 'float:15.2'),
                                          ('al2', 'float:15.2'),
                                          ('p_meh2', 'float:15.2'),
                                          ('p_rem2', 'float:15.2'),
                                          ('p_res2', 'float:15.2'),
                                          ('sat_k2', 'float:15.2'),
                                          ('k2', 'float:15.2'),
                                          ('rel_ca_mg2', 'float:15.2'),
                                          ('sat_mg2', 'float:15.2'),
                                          ('mg2', 'float:15.2'),
                                          ('sat_ca2', 'float:15.2'),
                                          ('ca2', 'float:15.2'),
                                          ('v2', 'float:15.2'),
                                          ('ph2', 'float:15.2'),
                                          ('ctc2', 'float:15.2'),
                                          ('mo2', 'float:15.2'),
                                          ('argila2', 'float:15.2')])}

    schema3 = {'geometry': 'Point',
               'properties': OrderedDict([('ID', 'int'),
                                          ('lab', 'str'),
                                          ('lote', 'str'),
                                          ('prof', 'str'),
                                          ('zn3', 'float:15.2'),
                                          ('mn3', 'float:15.2'),
                                          ('fe3', 'float:15.2'),
                                          ('cu3', 'float:15.2'),
                                          ('b3', 'float:15.2'),
                                          ('s3', 'float:15.2'),
                                          ('sat_al3', 'float:15.2'),
                                          ('al3', 'float:15.2'),
                                          ('p_meh3', 'float:15.2'),
                                          ('p_rem3', 'float:15.2'),
                                          ('p_res3', 'float:15.2'),
                                          ('sat_k3', 'float:15.2'),
                                          ('k3', 'float:15.2'),
                                          ('rel_ca_mg3', 'float:15.2'),
                                          ('sat_mg3', 'float:15.2'),
                                          ('mg3', 'float:15.2'),
                                          ('sat_ca3', 'float:15.2'),
                                          ('ca3', 'float:15.2'),
                                          ('v3', 'float:15.2'),
                                          ('ph3', 'float:15.2'),
                                          ('ctc3', 'float:15.2'),
                                          ('mo3', 'float:15.2'),
                                          ('argila3', 'float:15.2')])}

    schema4 = {'geometry': 'Point',
               'properties': OrderedDict([('ID', 'int'),
                                          ('lab', 'str'),
                                          ('lote', 'str'),
                                          ('prof', 'str'),
                                          ('zn4', 'float:15.2'),
                                          ('mn4', 'float:15.2'),
                                          ('fe4', 'float:15.2'),
                                          ('cu4', 'float:15.2'),
                                          ('b4', 'float:15.2'),
                                          ('s4', 'float:15.2'),
                                          ('sat_al4', 'float:15.2'),
                                          ('al4', 'float:15.2'),
                                          ('p_meh4', 'float:15.2'),
                                          ('p_rem4', 'float:15.2'),
                                          ('p_res4', 'float:15.2'),
                                          ('sat_k4', 'float:15.2'),
                                          ('k4', 'float:15.2'),
                                          ('rel_ca_mg4', 'float:15.2'),
                                          ('sat_mg4', 'float:15.2'),
                                          ('mg4', 'float:15.2'),
                                          ('sat_ca4', 'float:15.2'),
                                          ('ca4', 'float:15.2'),
                                          ('v4', 'float:15.2'),
                                          ('ph4', 'float:15.2'),
                                          ('ctc4', 'float:15.2'),
                                          ('mo4', 'float:15.2'),
                                          ('argila4', 'float:15.2')])}

    schema5 = {'geometry': 'Point',
               'properties': OrderedDict([('ID', 'int'),
                                          ('lab', 'str'),
                                          ('lote', 'str'),
                                          ('prof', 'str'),
                                          ('zn5', 'float:15.2'),
                                          ('mn5', 'float:15.2'),
                                          ('fe5', 'float:15.2'),
                                          ('cu5', 'float:15.2'),
                                          ('b5', 'float:15.2'),
                                          ('s5', 'float:15.2'),
                                          ('sat_al5', 'float:15.2'),
                                          ('al5', 'float:15.2'),
                                          ('p_meh5', 'float:15.2'),
                                          ('p_rem5', 'float:15.2'),
                                          ('p_res5', 'float:15.2'),
                                          ('sat_k5', 'float:15.2'),
                                          ('k5', 'float:15.2'),
                                          ('rel_ca_mg5', 'float:15.2'),
                                          ('sat_mg5', 'float:15.2'),
                                          ('mg5', 'float:15.2'),
                                          ('sat_ca5', 'float:15.2'),
                                          ('ca5', 'float:15.2'),
                                          ('v5', 'float:15.2'),
                                          ('ph5', 'float:15.2'),
                                          ('ctc5', 'float:15.2'),
                                          ('mo5', 'float:15.2'),
                                          ('argila5', 'float:15.2')])}

    schema6 = {'geometry': 'Point',
               'properties': OrderedDict([('ID', 'int'),
                                          ('lab', 'str'),
                                          ('lote', 'str'),
                                          ('prof', 'str'),
                                          ('zn6', 'float:15.2'),
                                          ('mn6', 'float:15.2'),
                                          ('fe6', 'float:15.2'),
                                          ('cu6', 'float:15.2'),
                                          ('b6', 'float:15.2'),
                                          ('s6', 'float:15.2'),
                                          ('sat_al6', 'float:15.2'),
                                          ('al6', 'float:15.2'),
                                          ('p_meh6', 'float:15.2'),
                                          ('p_rem6', 'float:15.2'),
                                          ('p_res6', 'float:15.2'),
                                          ('sat_k6', 'float:15.2'),
                                          ('k6', 'float:15.2'),
                                          ('rel_ca_mg6', 'float:15.2'),
                                          ('sat_mg6', 'float:15.2'),
                                          ('mg6', 'float:15.2'),
                                          ('sat_ca6', 'float:15.2'),
                                          ('ca6', 'float:15.2'),
                                          ('v6', 'float:15.2'),
                                          ('ph6', 'float:15.2'),
                                          ('ctc6', 'float:15.2'),
                                          ('mo6', 'float:15.2'),
                                          ('argila6', 'float:15.2')])}

    schema7 = {'geometry': 'Point',
               'properties': OrderedDict([('ID', 'int'),
                                          ('lab', 'str'),
                                          ('lote', 'str'),
                                          ('prof', 'str'),
                                          ('zn7', 'float:15.2'),
                                          ('mn7', 'float:15.2'),
                                          ('fe7', 'float:15.2'),
                                          ('cu7', 'float:15.2'),
                                          ('b7', 'float:15.2'),
                                          ('s7', 'float:15.2'),
                                          ('sat_al7', 'float:15.2'),
                                          ('al7', 'float:15.2'),
                                          ('p_meh7', 'float:15.2'),
                                          ('p_rem7', 'float:15.2'),
                                          ('p_res7', 'float:15.2'),
                                          ('sat_k7', 'float:15.2'),
                                          ('k7', 'float:15.2'),
                                          ('rel_ca_mg7', 'float:15.2'),
                                          ('sat_mg7', 'float:15.2'),
                                          ('mg7', 'float:15.2'),
                                          ('sat_ca7', 'float:15.2'),
                                          ('ca7', 'float:15.2'),
                                          ('v7', 'float:15.2'),
                                          ('ph7', 'float:15.2'),
                                          ('ctc7', 'float:15.2'),
                                          ('mo7', 'float:15.2'),
                                          ('argila7', 'float:15.2')])}

    for i in shp_list:
        try:
            if '0 A 20 CM' == i['prof'].head().iloc[0]:
                i.to_file(filename='download/0-20_' + et + '.shz', driver='ESRI Shapefile', schema=schema1,
                          compression="gzip")

        except:
            pass

        try:
            if '0 A 25 CM' == i['prof'].head().iloc[0]:
                i.to_file(filename='download/0-25_' + et + '.shz', driver='ESRI Shapefile', schema=schema1,
                          compression="gzip")

        except:
            pass

        try:
            if '0 A 10 CM' == i['prof'].head().iloc[0]:
                i.to_file(filename='download/0-10_' + et + '.shz', driver='ESRI Shapefile', schema=schema2)
        except:
            pass

        try:
            if '10 A 20 CM' == i['prof'].head().iloc[0]:
                i.to_file(filename='download/10-20_' + et + '.shz', driver='ESRI Shapefile', schema=schema3)
        except:
            pass

        try:
            if '20 A 40 CM' == i['prof'].head().iloc[0]:
                i.to_file(filename='download/20-40_' + et + '.shz', driver='ESRI Shapefile', schema=schema4)
        except:
            pass

        try:
            if '25 A 50 CM' == i['prof'].head().iloc[0]:
                i.to_file(filename='download/25-50_' + et + '.shz', driver='ESRI Shapefile', schema=schema4)
        except:
            pass

        try:
            if '40 A 60 CM' == i['prof'].head().iloc[0]:
                i.to_file(filename='download/40-60_' + et + '.shz', driver='ESRI Shapefile', schema=schema5)
        except:
            pass

        try:
            if '60 A 80 CM' == i['prof'].head().iloc[0]:
                i.to_file(filename='download/60-80_' + et + '.shz', driver='ESRI Shapefile', schema=schema6)
        except:
            pass

        try:
            if '80 A 100 CM' == i['prof'].head().iloc[0]:
                i.to_file(filename='download/80-100_' + et + '.shz', driver='ESRI Shapefile', schema=schema7)
        except:
            pass

    txt_list = []
    for i in shp_list:
        txt = i.drop(columns=['geometry'])
        txt_list.append(txt)

    for i in txt_list:
        try:
            if '0 A 20 CM' == i['prof'].head().iloc[0]:
                i.to_csv('download/0-20_' + et + '.txt', sep='\t', index=False)

        except:
            pass

        try:
            if '0 A 25 CM' == i['prof'].head().iloc[0]:
                i.to_csv('download/0-25_' + et + '.txt', sep='\t', index=False)

        except:
            pass

        try:
            if '0 A 10 CM' == i['prof'].head().iloc[0]:
                i.to_csv('download/0-10_' + et + '.txt', sep='\t', index=False)
        except:
            pass

        try:
            if '10 A 20 CM' == i['prof'].head().iloc[0]:
                i.to_csv('download/10-20_' + et + '.txt', sep='\t', index=False)
        except:
            pass

        try:
            if '20 A 40 CM' == i['prof'].head().iloc[0]:
                i.to_csv('download/20-40_' + et + '.txt', sep='\t', index=False)
        except:
            pass

        try:
            if '25 A 50 CM' == i['prof'].head().iloc[0]:
                i.to_csv('download/25-50_' + et + '.txt', sep='\t', index=False)
        except:
            pass

        try:
            if '40 A 60 CM' == i['prof'].head().iloc[0]:
                i.to_csv('download/40-60_' + et + '.txt', sep='\t', index=False)
        except:
            pass

        try:
            if '60 A 80 CM' == i['prof'].head().iloc[0]:
                i.to_csv('download/60-80_' + et + '.txt', sep='\t', index=False)
        except:
            pass

        try:
            if '80 A 100 CM' == i['prof'].head().iloc[0]:
                i.to_csv('download/80-100_' + et + '.txt', sep='\t', index=False)
        except:
            pass

def __relatorio_docx():

    import getpass
    from datetime import datetime
    from docx import Document
    from docx.shared import Inches


    USUARIO = getpass.getuser()
    DATA_HORA = datetime.now().strftime("%d-%m-%Y %H:%M:%S")

    document = Document('assets/MODELO.docx')
    # document.add_heading('Certificação Pré - Processamento', 0)
    tables = document.tables

    section = document.sections[0]
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.text = (USUARIO + '  ' + DATA_HORA)

    # GRAFICO JOIN
    GRAPH_JOIN_TABLE = tables[1].rows[0].cells[0].add_paragraph()
    r_1 = GRAPH_JOIN_TABLE.add_run()
    r_1.add_picture('download/graph_join.png', width=Inches(6))
    try:
        # RESPOSTA JOIN
        REPOSTA_JOIN_TABLE = tables[2].rows[1].cells[1].add_paragraph()
        r_2 = REPOSTA_JOIN_TABLE.add_run()
        r_2.text = (resposta_join)
    except:
        pass

    try:
        LAYOUT_GERAL = tables[4].rows[0].cells[0].add_paragraph()
        r_3 = LAYOUT_GERAL.add_run()
        r_3.add_picture('download/layout_map.png', width=Inches(8))
    except:
        pass

    try:
        # RESPOSTA ANALISE LAYOUT
        REPOSTA_LAYOUT_TABLE = tables[5].rows[0].cells[1].add_paragraph()
        r_4 = REPOSTA_LAYOUT_TABLE.add_run()
        r_4.text = (resposta_header_co + ' ; ' + resposta_overlap + ' ; ' + pontos_fora + ' ; ' + resposta_erro_text)

    except:
        pass
    try:
        # RESPOSTA ANALISE LAYOUT
        REPOSTA_DUPLICADO = tables[2].rows[0].cells[1].add_paragraph()
        r_6 = REPOSTA_DUPLICADO.add_run()
        r_6.text = (duplicate_str)

    except:
        pass

    try:
        # RESPOSTA LAYOUT
        REPOSTA_LAYOUT_TABLE = tables[5].rows[1].cells[1].add_paragraph()
        r_5 = REPOSTA_LAYOUT_TABLE.add_run()
        r_5.text = (resposta_layout)

    except:
        pass


    for png in list_png:
        document.add_picture('download/' + png, width=Inches(8))
        last_paragraph = document.paragraphs[-1]
        #last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        table2 = document.add_table(rows=1, cols=2)
        # Adding heading in the 1st row of the table
        row = table2.rows[0].cells
        table2.style = 'Table Grid'
        row[0].text = 'Observação usuário – Perfil de Amostras:'
        row[1].text = resposta_join_layout
    except:
        pass
    document.save("download/relatorio.docx")
    convert("download/relatorio.docx")
    trailer = PdfReader("download/relatorio.pdf")
    trailer.Info.WhoAmI = (USUARIO + '  ' + DATA_HORA)
    PdfWriter("download/relatorio.pdf", trailer=trailer).write()

def __download__():
    from zipfile import ZipFile

    # create a ZipFile object
    zipObj = ZipFile('relatorio.zip', 'w')
    # Add multiple files to the zip
    zipObj.write('download/estatistica_determinações.xlsx')
    zipObj.write('download/relatorio.pdf')
    zipObj.write('download/relatorio.docx')
    zipObj.write('download/without_join.xlsx')
    import os

    files = os.listdir('download/')
    for f in files:
        if '.shz' in f:
            zipObj.write('download/' + f)
    for f in files:
        if '.txt' in f:
            zipObj.write('download/' + f)

            # close the Zip File
    zipObj.close()


app = Dash(__name__, external_stylesheets=[dbc.themes.JOURNAL],
                  meta_tags=[{'name': 'viewport',
                              'content': 'width=device-width, initial-scale=1.0'}]
                  )



"""
        ORGANIZAÇÃO PARA APLICAÇÃO JUPYTER DASH
"""

# PLOT JOIN DF
fig = px.histogram(join_df, x="Profundidade", y="Nº de amostras",
                   color='Relação', barmode='group', text_auto='.4s',
                   # histfunc='avg',
                   height=500,
                   template="plotly_white")

# PLOT LAYOUT GERAL
fig2 = px.scatter_mapbox(point,
                         lat=point.geometry.y,
                         lon=point.geometry.x,
                         zoom=10,
                         hover_name="ID",
                         height=800,
                         width=1400
                         ).update_traces(marker={"size": 4}).update_layout(
    mapbox={
        "style": "open-street-map",
        "zoom": 11,
        "layers": [
            {
                "source": json.loads(farm_copy.geometry.to_json()),
                "below": "traces",
                "type": "line",
                "color": "orange",
                "line": {"width": 1.5},
            }
        ],
    })

if '!#ERROR#!' in duplicate_str:
    alerta_duplicado = dbc.Alert(duplicate_str, color="danger")
else:
    alerta_duplicado = ''

if '!#ERROR#!' in resposta_header_co:
    alerta_header = dbc.Alert(resposta_header_co, color="danger")
else:
    alerta_header = ''

if '!#ERROR#!' in resposta_erro_text:
    alerta_text = dbc.Alert(resposta_erro_text, color="danger")
else:
    alerta_text = ''

try:
    if '!#ERROR#!' in resposta_overlap:
        alerta_over = dbc.Alert(resposta_overlap, color="danger")
    else:
        alerta_over = ''
except:
    pass
if '!#ERROR#!' in pontos_fora:
    pontos_fora_alert = dbc.Alert(pontos_fora, color="danger")
else:
    pontos_fora_alert = ''


def table_type(df_column):
    # Note - this only works with Pandas >= 1.0.0

    if sys.version_info < (3, 0):  # Pandas 1.0.0 does not support Python 2
        return 'any'

    if isinstance(df_column.dtype, pd.DatetimeTZDtype):
        return 'datetime',
    elif (isinstance(df_column.dtype, pd.StringDtype) or
          isinstance(df_column.dtype, pd.BooleanDtype) or
          isinstance(df_column.dtype, pd.CategoricalDtype) or
          isinstance(df_column.dtype, pd.PeriodDtype)):
        return 'text'
    elif (isinstance(df_column.dtype, pd.SparseDtype) or
          isinstance(df_column.dtype, pd.IntervalDtype) or
          isinstance(df_column.dtype, pd.Int8Dtype) or
          isinstance(df_column.dtype, pd.Int16Dtype) or
          isinstance(df_column.dtype, pd.Int32Dtype) or
          isinstance(df_column.dtype, pd.Int64Dtype)):
        return 'numeric'
    else:
        return 'any'


app.layout = html.Div([

    dbc.Row([
        dbc.Col(html.H4("Análise Descritiva - Determinações",
                        className='text-center text-primary, mb-4'),
                width=12)

    ]),

    dbc.Row([
        dbc.Col([
            dcc.Dropdown(
                id="dropdown-id",
                options=[{"label": p, "value": p} for p in prof_list],
                placeholder="-Select a Profundity-",
                multi=False

            )], width={'size': 5})
    ]),

    dbc.Row([
        dbc.Col([

            # dbc.Table.from_dataframe(df, id='sts_df',striped=True, bordered=True, hover=True),
            dash_table.DataTable(df.to_dict('records'),
                                 [{"name": i, "id": i, 'type': table_type(df[i])} for i in df.columns],
                                 id='data-table-id', filter_action='native',
                                 style_table={'overflowX': 'auto', 'height': 390},
                                 style_cell={
                                     'height': 'auto',
                                     # all three widths are needed
                                     'minWidth': '50px', 'width': '50px', 'maxWidth': '50px',
                                     'whiteSpace': 'normal',
                                     'textAlign': 'center',
                                     'font_size': '12px'

                                 },
                                 style_data_conditional=[{
                                     'if': {'column_id': 'Máx',
                                            'filter_query': '{Máx} > {tolerancia}'},
                                     'backgroundColor': '#FF4136',
                                     'color': 'white',
                                 }],
                                 style_data={
                                     'width': '{}%'.format(10. / len(df.columns)),
                                     'textOverflow': 'hidden',
                                     'whiteSpace': 'normal',
                                     'height': 'auto',

                                 }, fill_width=False)]),
        dbc.Col([
            html.H6('Relação de Amostras duplicadas'),

            # dbc.Table.from_dataframe(df, id='sts_df',striped=True, bordered=True, hover=True),

            dash_table.DataTable(duplicate_df.to_dict('records'),
                                 [{"name": i, "id": i, 'type': table_type(duplicate_df[i])} for i in
                                  duplicate_df.columns], filter_action='native',
                                 style_table={'overflowX': 'auto', 'height': 390},
                                 style_cell={
                                     'height': 'auto',
                                     # all three widths are needed
                                     'minWidth': '50px', 'width': '50px', 'maxWidth': '50px',
                                     'whiteSpace': 'normal',
                                     'textAlign': 'center',
                                     'font_size': '12px'

                                 },

                                 style_data={
                                     'width': '{}%'.format(10. / len(duplicate_df.columns)),
                                     'textOverflow': 'hidden',
                                     'whiteSpace': 'normal',
                                     'height': 'auto',

                                 }, fill_width=False)

        ], align="center")]),

    dbc.Row([
        dbc.Col(html.H4("Relação Nº de Amostras e Amostragem (Join)",
                        className='text-center text-primary, mb-4'),
                width=12)]),
    dbc.Row([
    dbc.Row([alerta_duplicado]),
        dbc.Col([
            dbc.Label("Observação - Nº de Amostras e Join"),
            dbc.Input(id="input_join", placeholder="Type something...", type="text"),
            html.Br(),
            html.P(id="output_join")
        ], width={'size': 4})
    ]),
    dbc.Row([
        dbc.Row([
            dbc.Col([
                dbc.Button("Download Table", color="success", className="lg", id="btn_image_join",
                           style={
                               'display': 'inline-block',
                               'align': 'center',
                               'color': 'white', 'marginLeft': '5px',
                               'fontSize': '12px ',
                               'backgroundColor': '#101820',
                               'width': '130px',
                               'height': '40px',
                               'marginRight': '100px',
                               'MarginBottom': '40px'
                           })]),
            dcc.Download(id="download-file_join")]),

        dbc.Col([html.Div(
            # dbc.Table.from_dataframe(df, id='sts_df',striped=True, bordered=True, hover=True),
            dash_table.DataTable(without_join_df.to_dict('records'),
                                 [{"name": i, "id": i, 'type': table_type(without_join_df[i])} for i in
                                  without_join_df.columns], id='data-table-id_join', filter_action='native',
                                 style_table={'overflowX': 'auto', 'height': 740},
                                 style_cell={
                                     'height': 'auto',
                                     # all three widths are needed
                                     'minWidth': '50px', 'width': '50px', 'maxWidth': '50px',
                                     'whiteSpace': 'normal',
                                     'textAlign': 'center',
                                     'font_size': '12px'

                                 },
                                 style_data={
                                     'width': '{}%'.format(10. / len(without_join_df.columns)),
                                     'textOverflow': 'hidden',
                                     'whiteSpace': 'normal',
                                     'height': 'auto',

                                 }, fill_width=False))], width={"size": 5}),

        dbc.Col([html.Div(
            dcc.Graph(figure=fig, style={'width': '100vh', 'height': '80vh', 'size': 6}))
        ])

    ]),

    dbc.Row([
        dbc.Col(html.H4("Layout Geral",
                        className='text-center text-primary, mb-4'),
                width=12)]),
    dbc.Col([
        dbc.Label("Observação - Layout geral"),
        dbc.Input(id="input_layout", placeholder="Type something...", type="text"),
        html.Br(),
        html.P(id="output_layout")
    ], width={'size': 4}),
    dbc.Row([alerta_header]),
    dbc.Row([alerta_over]),
    dbc.Row([alerta_text]),
    dbc.Row([pontos_fora_alert]),
    dbc.Col([
        dcc.Graph(figure=fig2)
    ], width={'offset': 2}),
    dbc.Row([
        dbc.Col(html.H4("Perfil de Amostragem",
                        className='text-center text-primary, mb-4'),
                width=12)

    ]),
    dbc.Col([
        dbc.Label("Observação - Perfil de Amostragem"),
        dbc.Input(id="input_join_layout", placeholder="Type something...", type="text"),
        html.Br(),
        html.P(id="output_join_layout")
    ], width={'size': 6}),
    dbc.Row([
        dbc.Col([
            dcc.Dropdown(
                id="dropdown-id1",
                options=[{"label": p, "value": p} for p in prof_list],
                placeholder="-Select a Profundity-",
                multi=False)], width={'size': 5})
    ]),
    dbc.Col([
        dcc.Graph(id='graph-with-drop')], width={'offset': 2}),
    dbc.Col([
        dbc.Label("INSERIR ETAPA"),
        dbc.Input(id="input_etapa", placeholder="Type something...", type="text"),
        html.Br(),
        html.P(id="output_etapa")
    ], width={'size': 2}),
    dbc.Row([
        dbc.Row([
            dbc.Spinner(html.Div(id="loading-output"))]),
        dbc.Button("Download", color="success", className="me-1", id="btn_image"),
        dcc.Download(id="download-file")]),



])

### FUNÇÕES

@app.callback(
    Output("loading-output", "children"), [Input("btn_image", "n_clicks")]
)
def load_output(n):
    if n:
        time.sleep(1)
        return f"Output download {n} vezes"
    return "AO CLICAR, AGUARDE O DOWNLOAD - Ao fazer o download, certifiquesse se todos arquivos estão presentes no arquivo zipados"

##TABELA PROFUNDIDADE QUADRO ESTATISITCO

@app.callback(
    Output('data-table-id', 'data'),
    Input('dropdown-id', 'value')
)
def callback_func(dropdown_value):
    df_filtered = df[df['Prof'].eq(dropdown_value)]

    return df_filtered.to_dict(orient='records')

## CAIXA OBSERVAÇÃO JOIN
resposta_join = ''

@app.callback(Output("output_join", "children"), [Input("input_join", "value")])
def output_text_join(value):
    global resposta_join
    resposta_join = value
    return ' '

##DOWNLOAD TABLE
@app.callback(
    Output("download-file_join", "data"),
    Input("btn_image_join", "n_clicks"),
    prevent_initial_call=True,
)
def func_table(n_clicks):
    without_join_df.to_excel('without_join.xlsx', sheet_name='join', index=False)

    return dcc.send_file('without_join.xlsx')

## ## CAIXA OBSERVAÇÃO LAYOUT
resposta_layout = ''

@app.callback(Output("output_layout", "children"), [Input("input_layout", "value")])
def output_text_layout(value):
    global resposta_layout
    resposta_layout = value
    return ' '

resposta_join_layout = ''

@app.callback(Output("output_join_layout", "children"), [Input("input_join_layout", "value")])
def output_text_join_layout(value):
    global resposta_join_layout
    resposta_join_layout = value
    return ' '


et = ''

@app.callback(Output("output_etapa", "children"), [Input("input_etapa", "value")])
def output_etapa(value):
    global et
    et = value
    return ''

@app.callback(
    Output("download-file", "data"),
    Input("btn_image", "n_clicks"),
    prevent_initial_call=True,
)
def func(n_clicks):
    __relatorio_docx()
    __download_shp__(path, path_pt_shp)
    __download__()
    functions.__remove_files()
    return dcc.send_file('relatorio.zip'

                         )

@app.callback(
    Output('graph-with-drop','figure'),
    Input('dropdown-id1','value')
)
def update_figure(selected_prof):

    join_pt_plot = join_pt.loc[join_pt.prof == selected_prof].reset_index()

    fig = px.scatter_mapbox(join_pt_plot,
                        lat=join_pt_plot.geometry.y,
                        lon=join_pt_plot.geometry.x,
                        zoom=10,
                        hover_name="ID",
                        height = 800,
                        width = 1400,
                        title = selected_prof,

                        ).update_traces(marker={"size": 4}).update_layout(
    mapbox={
        "style": "open-street-map",
        "zoom": 11,
        "layers": [
            {
                "source": json.loads(farm_copy.geometry.to_json()),
                "below": "traces",
                "type": "line",
                "color": "orange",
                "line": {"width": 1.5},
            }
        ]
    })

    return fig


if __name__ == '__main__':
    Timer(1, open_browser).start();
    app.run_server(debug=False)




